""" import docx
import re
import matplotlib.pyplot as plt

# Read the content from the provided file
with open('daily_report_2023-10-30.txt', 'r') as file:
    content = file.read()

# Extract the date from the file name
match = re.search(r'\d{4}-\d{2}-\d{2}', 'daily_report_2023-10-30.txt')
date = match.group()

# Extract relevant information for graphs
# Black_ONT
uptime_black = float(re.search(r'Black_ONT Uptime: (\d+\.\d+) seconds', content).group(1))
downtime_black = float(re.search(r'Black_ONT Downtime: (\d+\.\d+) seconds', content).group(1))
total_time_black = float(re.search(r'Black_ONT Total Time: (\d+\.\d+) seconds', content).group(1))
availability_black = float(re.search(r'Black_ONT Network Availability: (\d+\.\d+)%', content).group(1))

# White_ONT
uptime_white = float(re.search(r'White_ONT Uptime: (\d+\.\d+) seconds', content).group(1))
downtime_white = float(re.search(r'White_ONT Downtime: (\d+\.\d+) seconds', content).group(1))
total_time_white = float(re.search(r'White_ONT Total Time: (\d+\.\d+) seconds', content).group(1))
availability_white = float(re.search(r'White_ONT Network Availability: (\d+\.\d+)%', content).group(1))

# Create individual pie charts for Uptime, Downtime, and Total Time
# Black_ONT
labels = ['Uptime', 'Downtime']
values = [uptime_black, downtime_black]

fig, ax = plt.subplots()
ax.pie(values, labels=labels, autopct='%1.1f%%')
ax.set(aspect="equal")
plt.title('Black_ONT Network Status')

plt.savefig('black_ont_pie_chart.png')
plt.close()

# White_ONT
values = [uptime_white, downtime_white]

fig, ax = plt.subplots()
ax.pie(values, labels=labels, autopct='%1.1f%%')
ax.set(aspect="equal")
plt.title('White_ONT Network Status')

plt.savefig('white_ont_pie_chart.png')
plt.close()

# Create individual bar graphs for Network Availability
# Black_ONT
plt.bar('Black_ONT', availability_black, color='green')
plt.xlabel('Devices')
plt.ylabel('Availability (%)')
plt.title('Black_ONT Network Availability')

plt.savefig('black_ont_bar_chart.png')
plt.close()

# White_ONT
plt.bar('White_ONT', availability_white, color='green')
plt.xlabel('Devices')
plt.ylabel('Availability (%)')
plt.title('White_ONT Network Availability')

plt.savefig('white_ont_bar_chart.png')
plt.close()

# Create a Word document with charts and information for each ONT
doc = docx.Document()
doc.add_heading(f"Network Availability until {date}", level=1)

# Black_ONT
doc.add_paragraph('Black_ONT Uptime, Downtime, and Total Time (Pie Chart):')
doc.add_picture('black_ont_pie_chart.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

doc.add_paragraph(f"Black_ONT Network Availability: {availability_black}%")
doc.add_picture('black_ont_bar_chart.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

# White_ONT
doc.add_paragraph('White_ONT Uptime, Downtime, and Total Time (Pie Chart):')
doc.add_picture('white_ont_pie_chart.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

doc.add_paragraph(f"White_ONT Network Availability: {availability_white}%")
doc.add_picture('white_ont_bar_chart.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

doc.save(f"Network_Availability_{date}.docx") """




import docx
import re
import matplotlib.pyplot as plt
import sys

def generate_reports(frequency, file_name):
    # Read the content from the provided file
    with open(file_name, 'r') as file:
        content = file.read()

    # Extract the date from the file name
    match = re.search(r'\d{4}-\d{2}-\d{2}', file_name)
    date = match.group()

    # Extract relevant information for graphs
    # Black_ONT
    uptime_black = float(re.search(r'Black_ONT Uptime: (\d+\.\d+) seconds', content).group(1))
    downtime_black = re.search(r'Black_ONT Downtime: (\d+\.\d+) seconds', content)
    if downtime_black:
        downtime_black = float(downtime_black.group(1))
    else:
        # Handle the case where the pattern isn't found in the file
        downtime_black = 0  # Set a default value or handle it based on your requirements
    total_time_black = float(re.search(r'Black_ONT Total Time: (\d+\.\d+) seconds', content).group(1))
    availability_black = float(re.search(r'Black_ONT Network Availability: (\d+\.\d+)%', content).group(1))

    # White_ONT
    uptime_white = float(re.search(r'White_ONT Uptime: (\d+\.\d+) seconds', content).group(1))
    downtime_white = re.search(r'White_ONT Downtime: (\d+\.\d+) seconds', content)
    if downtime_white:
        downtime_white = float(downtime_white.group(1))
    else:
        # Handle the case where the pattern isn't found in the file
        downtime_white = 0  # Set a default value or handle it based on your requirements
    total_time_white = float(re.search(r'White_ONT Total Time: (\d+\.\d+) seconds', content).group(1))
    availability_white = float(re.search(r'White_ONT Network Availability: (\d+\.\d+)%', content).group(1))

    # Create individual pie charts for Uptime, Downtime, and Total Time
    # Black_ONT
    labels = ['Uptime', 'Downtime']
    values = [uptime_black, downtime_black]

    fig, ax = plt.subplots()
    ax.pie(values, labels=labels, autopct='%1.1f%%')
    ax.set(aspect="equal")
    plt.title('Black_ONT Network Status')

    plt.savefig('black_ont_pie_chart.png')
    plt.close()

    # White_ONT
    values = [uptime_white, downtime_white]

    fig, ax = plt.subplots()
    ax.pie(values, labels=labels, autopct='%1.1f%%')
    ax.set(aspect="equal")
    plt.title('White_ONT Network Status')

    plt.savefig('white_ont_pie_chart.png')
    plt.close()

    # Create individual bar graphs for Network Availability
    # Black_ONT
    plt.bar('Black_ONT', availability_black, color='green')
    plt.xlabel('Devices')
    plt.ylabel('Availability (%)')
    plt.title('Black_ONT Network Availability')

    plt.savefig('black_ont_bar_chart.png')
    plt.close()

    # White_ONT
    plt.bar('White_ONT', availability_white, color='green')
    plt.xlabel('Devices')
    plt.ylabel('Availability (%)')
    plt.title('White_ONT Network Availability')

    plt.savefig('white_ont_bar_chart.png')
    plt.close()

    # Create a Word document with charts and information for each ONT
    doc = docx.Document()
    doc.add_heading(f"Network Availability until {date}", level=1)

    # Black_ONT
    doc.add_paragraph('Black_ONT Uptime, Downtime, and Total Time (Pie Chart):')
    doc.add_picture('black_ont_pie_chart.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

    doc.add_paragraph(f"Black_ONT Network Availability: {availability_black}%")
    doc.add_picture('black_ont_bar_chart.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

    # White_ONT
    doc.add_paragraph('White_ONT Uptime, Downtime, and Total Time (Pie Chart):')
    doc.add_picture('white_ont_pie_chart.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

    doc.add_paragraph(f"White_ONT Network Availability: {availability_white}%")
    doc.add_picture('white_ont_bar_chart.png', width=docx.shared.Inches(5), height=docx.shared.Inches(4))

    doc.save(f"Network_Availability_{frequency}_{date}.docx")

# Check for the argument specifying the report frequency and the file name
if len(sys.argv) != 3:
    print("Usage: create_graph_reports.py <frequency> <file_name>")
    sys.exit(1)

frequency = sys.argv[1]
file_name = sys.argv[2]

if frequency == 'daily' or frequency == 'weekly' or frequency == 'monthly':
    generate_reports(frequency, file_name)
else:
    print("Invalid frequency provided. Use 'daily', 'weekly', or 'monthly'.")
